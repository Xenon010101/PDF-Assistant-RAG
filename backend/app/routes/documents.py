"""
Document management routes — upload, list, delete, and serve PDF files.
Background ingestion via FastAPI BackgroundTasks.
"""
import os
import sys
import uuid
import logging
import asyncio
import concurrent.futures
from datetime import datetime, timezone
from typing import Optional
from pathlib import Path
import shutil
import tempfile
from urllib.parse import urlparse
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, BackgroundTasks, status, Query
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session

from app.database import get_db
from app.models import User, Document
from app.schemas import (
    DocumentResponse,
    DocumentListResponse,
    DocumentStatusResponse,
    DocumentRename,
    ChunkSettings,
    UploadUrl,
)
from app.auth import get_current_user
from app.config import get_settings
from app.rag.chunker import chunk_document, get_page_count
from app.rag.vectorstore import store_chunks

try:
    from crawl4ai import AsyncWebCrawler
    from crawl4ai.async_configs import BrowserConfig, CrawlerRunConfig
except ImportError as exc:
    AsyncWebCrawler = None
    BrowserConfig = None
    CrawlerRunConfig = None
    CRAWL4AI_IMPORT_ERROR = exc
else:
    CRAWL4AI_IMPORT_ERROR = None

from sqlalchemy import select
logger = logging.getLogger(__name__)
settings = get_settings()

router = APIRouter(prefix="/documents", tags=["Documents"])


ALLOWED_MIME_TYPES = settings.ALLOWED_MIME_TYPES


async def validate_upload(file: UploadFile):
    """Validate an uploaded file and save it to a temporary file.

    Checks the file extension, size (against `settings.MAX_UPLOAD_SIZE_MB`),
    MIME type via libmagic, and performs deep validation by attempting to
    parse the file (PDF with pypdf, DOCX with python-docx). On success,
    returns the path to a temporary saved file. 

    Args:
        file: The FastAPI UploadFile object to validate.

    Returns:
        str: Path to the temporary saved file that passed all validations.
    
    Raises:
        HTTPException: With status code 400 if:
            - No filename is provided.
            - The file extension is not allowed. (only .pdf or .docx)
            - The file size exceeds the maximum limit.
            - The MIME type does not match allowed types for the extension.
            - The file is corrupted or cannot be parsed. (invalid PDF/DOCX)
        HTTPException: With status code 500 if:
            - 'python-magic' dependency is missing on the server.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")

    ext = Path(file.filename).suffix.lower()

    # extension without leading dot in settings
    if ext.lstrip(".") not in settings.ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Only PDF, DOCX, TEXT, AND MARKDOWN files are allowed")

    # save to a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        # UploadFile.file is a file-like object
        shutil.copyfileobj(file.file, tmp)
        temp_path = tmp.name

    try:
        size = Path(temp_path).stat().st_size

        if size > settings.MAX_UPLOAD_SIZE_MB * 1024 * 1024:
            Path(temp_path).unlink(missing_ok=True)
            raise HTTPException(status_code=400, detail="File too large")

        # libmagic may not be installed in all environments — import lazily
        try:
            import magic
            # make sure you have installed libmagic in your system, otherwise it will not work
        except Exception:
            Path(temp_path).unlink(missing_ok=True)
            raise HTTPException(status_code=500, detail="Server missing 'python-magic' dependency")

        mime = magic.from_file(temp_path, mime=True)

        if mime not in ALLOWED_MIME_TYPES.get(ext, []):
            Path(temp_path).unlink(missing_ok=True)
            raise HTTPException(status_code=400, detail=f"Invalid file type: {mime}")

        # Deep validation: try to parse the file — import parsers lazily
        try:
            if ext == ".pdf":
                from pypdf import PdfReader

                PdfReader(temp_path)
            elif ext == ".docx":
                from docx import Document as DocxDocument

                DocxDocument(temp_path)
        except Exception:
            Path(temp_path).unlink(missing_ok=True)
            raise HTTPException(status_code=400, detail="Corrupted or invalid file")

        return temp_path

    finally:
        """If an exception was raised above it will propagate; caller should
        remove the temp file when appropriate. We don't unlink here because
        caller will move the file on success."""
        pass


def _ingest_document(document_id: str, filepath: str, original_name: str, user_id: str):
    """
    Process a document in the background: chunk document, generate embeddings, and store in ChromaDB,
    calls document summary function, and update the database record.

    This function is intended to be run as a background task. 
    It creates its own database session, updates the
    document status, extracts text, splits into chunks, generates embeddings,
    stores everything in ChromaDB, calls summary function, updates the document record with page count,
    chunk count, and summary, and marks the document as 'ready'. 
    On failure, it sets status to 'failed' and records the error message.

    Args:
        document_id: Unique identifier of the document in the database.
        filepath: Absolute or relative path to the uploaded file on disk.
        original_name: original filename provided by the user (for logging and metadata).
        user_id: Identifier of the user who owns the document.
    
    Returns:
        None

    Note:
        This function does not raise exceptions to the caller;
        all errors are logged and the database record is updated accordingly. 
    """
    from app.database import SessionLocal

    db = SessionLocal()
    try:
        doc = (
            db.query(Document)
            .filter(Document.id == document_id, Document.is_deleted.is_(False))
            .first()
        )
        if not doc:
            logger.error(f"Document {document_id} not found for ingestion")
            return

        # Update status to processing
        doc.status = "processing"
        db.commit()

        # Get page count
        page_count = get_page_count(filepath)
        doc.page_count = page_count

        # Chunk document with optional chunk size and overlap parameters from the document record, falling back to global defaults if not set
        chunk_size = doc.chunk_size
        chunk_overlap = doc.chunk_overlap
        try:
            kwargs = {}
            if chunk_size is not None:
                kwargs["chunk_size"] = chunk_size
            if chunk_overlap is not None:
                kwargs["chunk_overlap"] = chunk_overlap

            if kwargs:
                chunks = chunk_document(filepath, **kwargs)
            else:
                chunks = chunk_document(filepath)

        except TypeError:
            # Backward-compatible fallback for chunk_document implementations/tests
            # that only accept (filepath)
            chunks = chunk_document(filepath)

        if not chunks:
            doc.status = "failed"
            doc.error_message = "No text could be extracted from the document"
            db.commit()
            return

        # Build and persist a lightweight entity co-occurrence graph for GraphRAG.
        try:
            from app.rag.graph_builder import build_graph, save_graph

            graph = build_graph(chunks)
            save_graph(graph, user_id=user_id, document_id=document_id)
        except Exception as e:
            logger.warning(f"Could not build knowledge graph for document {document_id}: {e}")

        # Store embeddings in ChromaDB
        chunk_count = store_chunks(
            chunks=chunks,
            document_id=document_id,
            filename=original_name,
            user_id=user_id,
        )

        # Generate summary and update document record
        try:
            from app.rag.summarizer import generate_document_summary

            summary = generate_document_summary(filepath, max_sentences=2)
            if summary:
                doc.summary = summary
                db.commit() # Update document record with summary                
        except Exception as e:
            logger.warning(f"Could not import summarizer for document {document_id}: {e}")
            doc.summary = None

        # Update document record
        doc.chunk_count = chunk_count
        doc.status = "ready"
        db.commit()

        logger.info(f"Document {document_id} ingested: {page_count} pages, {chunk_count} chunks")

    except Exception as e:
        logger.error(f"Ingestion error for {document_id}: {e}")
        try:
            doc = (
                db.query(Document)
                .filter(Document.id == document_id, Document.is_deleted.is_(False))
                .first()
            )
            if doc:
                doc.status = "failed"
                doc.error_message = str(e)[:500]
                db.commit()
        except Exception:
            pass
    finally:
        db.close()



def _crawl_in_new_loop(url: str) -> str:
    """Run the async crawler in a fresh event loop on a worker thread.
    On Windows this must be a ProactorEventLoop to support subprocesses.
    """
    if sys.platform == "win32":
        loop = asyncio.ProactorEventLoop()
    else:
        loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    try:
        async def _crawl():
            browser_config = BrowserConfig()
            run_config = CrawlerRunConfig(
                excluded_tags=['form', 'header'],

                # Content processing
                process_iframes=True,
                # remove_overlay_elements=True,

                # Cache control
                # cache_mode=CacheMode.ENABLED
            )
            async with AsyncWebCrawler(config=browser_config) as crawler:
                result = await crawler.arun(url=url, config=run_config)
                return result.markdown or ""
        return loop.run_until_complete(_crawl())
    finally:
        loop.close()

@router.post("/upload", response_model=DocumentResponse, status_code=status.HTTP_202_ACCEPTED)
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """
    Upload a document and enqueue RAG processing.
    
    Validates the uploaded file (extension, size, MIME type, integrity),
    saves it to the user's directory, creates a database record with status
    'pending', schedules a background task for chunking and embedding, and
    returns 202 Accepted immediately so large documents do not block the API
    request while embeddings are generated.

    Args:
        background_tasks: FastAPI BackgroundTasks instance to run the ingestion process asynchronously.
        file: The uploaded file, provided as a multipart/form-data field in the request.
        user: The currently authenticated user, injected by the `get_current_user` dependency.
        db: Database session, injected by the `get_db` dependency.

    Returns:
        DocumentResponse: The created document record, validated against the
        response model (includes id, filename, original_name, file_size, status, etc.).

    Raises:
        HTTPException: With status code 400 if:
            - No filename is provided.
            - The file extension is not allowed. (only .pdf or .docx)
            - The file fails validation checks (size, MIME type, integrity).
        HTTPException: With status code 500 if:
            - The server lacks the 'python-magic' dependency. 
    """
    # ── Validate file type ───────────────────────────
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")

    ext = file.filename.rsplit(".", 1)[-1].lower()
    if ext not in settings.ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"File type '.{ext}' not supported. Allowed: {', '.join(settings.ALLOWED_EXTENSIONS)}",
        )

    # ── Validate and save file to disk ───────────────
    temp_path = await validate_upload(file)

    user_dir = os.path.join(settings.UPLOAD_DIR, user.id)
    os.makedirs(user_dir, exist_ok=True)

    stored_filename = f"{uuid.uuid4().hex}.{ext}"
    filepath = os.path.join(user_dir, stored_filename)

    # Move temp file to final destination
    shutil.move(temp_path, filepath)

    file_size = Path(filepath).stat().st_size

    # ── Create database record ───────────────────────
    document = Document(
        user_id=user.id,
        filename=stored_filename,
        original_name=file.filename,
        file_size=file_size,
        status="pending",
    )
    db.add(document)
    db.commit()
    db.refresh(document)

    # ── Trigger background ingestion ─────────────────
    background_tasks.add_task(
        _ingest_document,
        document_id=document.id,
        filepath=filepath,
        original_name=file.filename,
        user_id=user.id,
    )

    return DocumentResponse.model_validate(document)

@router.post("/urlupload", status_code=status.HTTP_202_ACCEPTED)
async def upload_document_url(
        payload: UploadUrl,
        background_tasks: BackgroundTasks,
        user: User = Depends(get_current_user),
        db: Session = Depends(get_db),
):
    """
    Uses crawl4ai's AsyncWebCrawler in a dedicated thread with its own
    event loop. This is required on Windows because uvicorn's default
    SelectorEventLoop does not support subprocess creation (used by
    Playwright/crawl4ai), which causes a NotImplementedError.
    On Linux (production) a plain new_event_loop() is used instead.
    """
    if CRAWL4AI_IMPORT_ERROR is not None:
        raise HTTPException(
            status_code=503,
            detail="URL upload is unavailable because crawl4ai is not installed",
        )

    temp_path: Optional[str] = None
    try:
        parsed = urlparse(payload.url)
        if not all([parsed.scheme, parsed.netloc]):
            raise HTTPException(status_code=400, detail="Invalid URL")


        # Run in a worker thread with its own event loop to avoid
        # NotImplementedError on Windows (SelectorEventLoop can't spawn subprocesses)
        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
            markdown = await asyncio.get_event_loop().run_in_executor(
                pool, _crawl_in_new_loop, payload.url
            )

        if not markdown:
            raise HTTPException(status_code=422, detail="No content could be extracted from the URL")


        with tempfile.NamedTemporaryFile(
            mode="w",
            suffix=".txt",
            delete=False,
            encoding="utf-8",
        ) as tmp:
            tmp.write(markdown)
            temp_path = tmp.name

        # ── Move temp file to permanent user upload directory ──
        ext = "txt"
        user_dir = os.path.join(settings.UPLOAD_DIR, user.id)
        os.makedirs(user_dir, exist_ok=True)

        stored_filename = f"{uuid.uuid4().hex}.{ext}"
        filepath = os.path.join(user_dir, stored_filename)
        shutil.move(temp_path, filepath)
        temp_path = None  # file is now at filepath; no longer a temp to clean up

        file_size = Path(filepath).stat().st_size

        # ── Derive a human-readable name from the URL ─────────
        url_path = parsed.path.rstrip("/")
        original_name = f"{parsed.netloc}{url_path or ''}.txt"

        # ── Create database record ─────────────────────────────
        document = Document(
            user_id=user.id,
            filename=stored_filename,
            original_name=original_name,
            file_size=file_size,
            status="pending",
        )
        db.add(document)
        db.commit()
        db.refresh(document)

        # ── Trigger background ingestion ───────────────────────
        background_tasks.add_task(
            _ingest_document,
            document_id=document.id,
            filepath=filepath,
            original_name=original_name,
            user_id=user.id,
        )

        return DocumentResponse.model_validate(document)

    except HTTPException:
        raise
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid URL")
    except Exception as e:
        logger.error(f"URL upload error: {e}")
        raise HTTPException(status_code=400, detail=f"Something went wrong with URL processing: {str(e)}")
    finally:
        '''Runs whether the request succeeded, raised an HTTPException,
        or hit an unexpected error — no temp files are ever left behind.'''
        if temp_path is not None:
            Path(temp_path).unlink(missing_ok=True)



@router.get("/{document_id}/status", response_model=DocumentStatusResponse)
def get_document_status(
    document_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """
    Poll processing status for a single uploaded document.

    This endpoint lets clients refresh the upload lifecycle without fetching
    the entire document list. The returned status is one of the existing
    document states: pending, processing, ready, or failed.
    """
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == user.id,
        Document.is_deleted.is_(False),
    ).first()

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    return DocumentStatusResponse.model_validate(doc)


@router.get("/", response_model=DocumentListResponse)
def list_documents(
    page: int = Query(1, ge=1),
    per_page: int = Query(20, ge=1),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """
    List all documents for the authenticated user with pagination.

    Returns a paginated list of documents belonging to the current user,
    ordered by upload date (newest first).

    Args:
        page: The page number to retrieve (1: indexed). Defaults to 1.
        per_page: The number of documents to return per page. Defaults to 20.
        user: The currently authenticated user, injected by the `get_current_user` dependency.
        db: Database session, injected by the `get_db` dependency.
        
    Returns:
        DocumentListResponse: A response model containing:
            - items: A list of DocumentResponse objects for the current page.
            - total: The total number of documents for the user.
            - page: The current page number.
            - pages: The total number of pages available.
    """

    """Number of rows to skip"""
    skip: int = (page - 1) * per_page

    """Total Pages"""
    totalDocuments = (
        db.query(Document)
        .filter(Document.user_id == user.id, Document.is_deleted.is_(False))
        .count()
    )
    """Total Pages"""
    pages = (totalDocuments + per_page - 1) // per_page
    
    """List all documents for the authenticated user in Paginated form"""
    docs = ((
            db.execute(select(Document)
            .where(Document.user_id == user.id, Document.is_deleted.is_(False))
            .order_by(Document.uploaded_at.desc())
            .limit(per_page).offset(skip))
            )
            .scalars().all())

    return DocumentListResponse(
        items=[DocumentResponse.model_validate(d) for d in docs],
        total=totalDocuments,
        page=page,
        pages=pages
    )


@router.patch("/{document_id}", response_model=DocumentResponse)
def rename_document(
    document_id: str,
    rename: DocumentRename,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """
    Rename an uploaded document without changing its stored file or vector data.
    """
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.is_deleted.is_(False),
    ).first()

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if str(doc.user_id) != str(user.id):
        raise HTTPException(status_code=403, detail="You do not have permission to rename this document")

    doc.original_name = rename.name
    db.commit()
    db.refresh(doc)

    return DocumentResponse.model_validate(doc)


@router.get("/{document_id}", response_model=DocumentResponse)
def get_document(
    document_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """
    Retrieve a specific document by its ID for the authenticated user.

    Fetches the document that matches both the provided `document_id` and
    the current user's ID. If no such document exists, a 404 error is raised.

    Args:
        document_id: The unique identifier of the document to retrieve.
        user: The currently authenticated user, injected by the `get_current_user` dependency.
        db: Database session, injected by the `get_db` dependency.

    Returns:
        DocumentResponse: The document record that matches the criteria, validated against the response model
        (includes id, filename, original_name, file_size, status, etc.).

    Raises:
        HTTPException: With status code 404 if the document is not found or does not belong to the authenticated user.
    """
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == user.id,
        Document.is_deleted.is_(False),
    ).first()

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    return DocumentResponse.model_validate(doc)


@router.get("/{document_id}/pdf")
def serve_pdf(
    document_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """
    Serve the PDF file for the document viewer.

    Retrieves the document from the database to verify ownership, then
    returns the actual PDF file from disk as a downloadable response.

    Args:
        document_id: The unique identifier of the document whose PDF is to be served.
        user: The currently authenticated user, injected by the `get_current_user` dependency.
        db: Database session, injected by the `get_db` dependency.

    Returns:
        FileResponse: A FastAPI FileResponse object that streams the PDF
        file to the client with the correct media type and original filename.
    
    Raises:
        HTTPException: 404 if the document does not exist or does not belong
            to the authenticated user, or if the file is missing on disk.
    """
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == user.id,
        Document.is_deleted.is_(False),
    ).first()

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    filepath = os.path.join(settings.UPLOAD_DIR, user.id, doc.filename)

    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="File not found on disk")

    return FileResponse(
        filepath,
        media_type="application/pdf",
        filename=doc.original_name,
    )


@router.delete("/{document_id}", status_code=status.HTTP_200_OK)
def delete_document(
    document_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """
    Soft-delete a document so it disappears from normal document APIs.

    The underlying file, vectors, graph, and chat history are retained for a
    future recycle-bin/restore flow. Normal read/list endpoints filter deleted
    documents so accidental deletion is reversible at the database level.

    Args:
        document_id: The unique identifier of the document to delete.
        user: The currently authenticated user, injected by the `get_current_user` dependency.
        db: Database session, injected by the `get_db` dependency.

    Returns:
        dict: A JSON response containing a success message confirming the deletion of the document.

    Raises:
        HTTPException: With status code 404 if the document is not found or does not belong to the authenticated user.

    Note:
        ChromaDB deletion errors are caught and logged only; they do not
        raise an HTTP exception because the main document record is already
        removed from the database.
    """
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == user.id,
        Document.is_deleted.is_(False),
    ).first()

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    doc.is_deleted = True
    doc.deleted_at = datetime.now(timezone.utc)
    db.commit()

    return {"message": f"Document '{doc.original_name}' deleted successfully"}


@router.post("/{document_id}/chunk_settings", response_model=DocumentResponse)
def update_chunk_settings(
    document_id: str,
    settings_update: ChunkSettings,
    background_tasks: BackgroundTasks,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Update chunking settings for a specific document.

    This endpoint allows users to update the chunk size and overlap for a document and calls _injest_document fucntion in the background to re-chunk the document with new chunk parameters.

    Args:
        document_id: The unique identifier of the document to update.
        settings_update: A ChunkSettings object containing the chunk_size and chunk_overlap values.
        background_tasks: FastAPI BackgroundTasks instance to run the ingestion process asynchronously.
        user: The currently authenticated user, injected by the `get_current_user` dependency.
        db: Database session, injected by the `get_db` dependency.

    Returns:
        DocumentResponse: The updated document record, validated against the response model.

    Raises:
        HTTPException: With status code 404 if the document is not found or does not belong to the authenticated user.
        HTTPException: With status code 400 if the provided chunk size or overlap values are invalid (e.g., chunk size less than 100, or overlap greater than or equal to chunk size).
    """
    # Validate if the document exists and belongs to the user
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == user.id,
        Document.is_deleted.is_(False),
    ).first()

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if settings_update.chunk_size is not None:
        if settings_update.chunk_size < 100:
            raise HTTPException(400, "Chunk size must be at least 100")
        doc.chunk_size = settings_update.chunk_size
    if settings_update.chunk_overlap is not None:
        if settings_update.chunk_overlap >= settings_update.chunk_size:
            raise HTTPException(400, "Chunk overlap cannot be greater than or equal to chunk size")
        doc.chunk_overlap = settings_update.chunk_overlap    

    # Refresh the document record to update the chunk settings before re-ingestion
    db.commit()
    db.refresh(doc)

    # Reset document status, chunk/page counts, summary to trigger re-ingestion with new chunk settings.
    doc.status = "pending"
    doc.chunk_count = 0
    doc.page_count = 0
    doc.summary = None
    db.commit()

    # Trigger background ingestion with updated chunk settings. The _ingest_document function will read the new chunk settings from the document record and re-chunk the document accordingly.
    background_tasks.add_task(
        _ingest_document,
        document_id=doc.id,
        filepath=os.path.join(settings.UPLOAD_DIR, user.id, doc.filename), 
        original_name=doc.original_name,
        user_id=user.id,
    )
    # Return the updated document record with new chunk settings
    return DocumentResponse.model_validate(doc)
